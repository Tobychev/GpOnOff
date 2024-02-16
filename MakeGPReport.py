#!/usr/bin/env python
import argparse
import astropy.table as at
import docx
import numpy as np
import os
formatters = {
      "counts_off": "{:>5d}".format,
      "counts": "{:>4d}".format,
      "livetime": lambda x: "{:.3f}".format(x/3600),
      "ontime": lambda x: "{:.3f}".format(x/3600),
      "stat_type": str,
       }

def make_report(loc):
   oldbinpic = rebinpic = None
   cbeini = cberes = cont = None

   print(f"Making report for {loc}")
   for dir in os.listdir(loc):
      breakpoint()
      if dir.endswith("ecsv"):
         if dir.endswith("fit_result.ecsv"):
            fit = at.Table.read(f"{loc}/{dir}",format="ascii.ecsv")
         if dir.endswith("cumul.ecsv"):
            cml = at.Table.read(f"{loc}/{dir}",format="ascii.ecsv")
         if dir.endswith("byrun.ecsv"):
            run = at.Table.read(f"{loc}/{dir}",format="ascii.ecsv")
      elif dir.endswith("png"):
         if dir.endswith("Stat.png"):
            statpic = dir
         if dir.endswith("Fit.png"):
            fitpic = dir
         if dir.endswith("Residuals.png"):
            respic = dir
         if dir.endswith("Likelihood.png"):
            likpic = dir
         if dir.endswith("Raw_Binning.png"):
            oldbinpic = dir
         if dir.endswith("Rebinned.png"):
            rebinpic = dir
         if dir.endswith("CountsByExpInitial.png"):
            cbeini = dir
         if dir.endswith("CountsByExpResampled.png"):
            cberes = dir
         if dir.endswith("Countours.png"):
            cont = dir
      elif dir.endswith("yaml"):
         conf = dir
   if not "statpic" in locals():
      raise ValueError(f"Folder {loc} is not a proper result location")
   if not "respic" in locals():
      raise ValueError(f"Folder {loc} is lacks a residuals plot, rerun OnOffAnalysis to fix this")

   doc = docx.Document()
   doc.add_heading(f"Source {loc}",0)
   doc.add_heading("Stats",1)
   doc.add_picture(f"{loc}/{statpic}",width=docx.shared.Cm(4))

   if cml:
      zens=  cml["zenith"]
      zen_mean = zens.mean()
      zen_med,zen_sty = np.percentile(zens,[50,75])
      tot_stat = cml[-1]
      nrows = len(tot_stat.colnames)+2
      #The above leads to euse the by run zenit spot
      table1 = doc.add_table(rows=nrows,cols=2)

      for idx,val in enumerate(tot_stat):
         name = tot_stat.colnames[idx]
         row = table1.rows[idx]
         if name == "name":
            row.cells[0].text = "Name"
            row.cells[1].text = "Value"
            continue
         frmt = formatters.get(name,"{:8.4f}".format)
         print(name,frmt(val))
         row.cells[0].text = name
         row.cells[1].text = str(frmt(val))
      row_med = table1.rows[-3]
      row_med.cells[0].text = "Median Zenith"
      row_med.cells[1].text = f"{zen_med:.3f}"
      row_mean = table1.rows[-2]
      row_mean.cells[0].text = "Mean Zenith"
      row_mean.cells[1].text = f"{zen_mean:.3f}"
      row_sty = table1.rows[-1]
      row_sty.cells[0].text = "Zenith 75%"
      row_sty.cells[1].text = f"{zen_sty:.3f}"
   else:
      raise ValueError(f"Did not find a stats file in {loc}")

   doc.add_heading("Spectral Fit",1)
   doc.add_picture(f"{loc}/{fitpic}",width=docx.shared.Cm(4))


   table2 = doc.add_table(rows=6,cols=4)
   row = table2.rows[0]
   row.cells[0].text = "Name"
   row.cells[1].text = "Value"
   row.cells[2].text = "Error"
   row.cells[3].text = "Unit"

   for idx,vals in enumerate(fit): 
      row = table2.rows[idx]
      name = vals[0]
      frmt = formatters.get(name,"{:.4e}".format)
      row.cells[0].text = name
      row.cells[1].text = frmt(vals[1])
      row.cells[2].text = frmt(vals[2])
      row.cells[3].text = str(vals[3])

   doc.add_heading("Diagnostics",2)
   if oldbinpic:
      doc.add_picture(f"{loc}/{oldbinpic}",width=docx.shared.Cm(4))
   if cbeini:
      doc.add_picture(f"{loc}/{cbeini}",width=docx.shared.Cm(4))
   if rebinpic:
      doc.add_picture(f"{loc}/{rebinpic}",width=docx.shared.Cm(4))
      doc.add_picture(f"{loc}/{cberes}",width=docx.shared.Cm(4))
   if cont:
      doc.add_picture(f"{loc}/{cont}",width=docx.shared.Cm(4))

   doc.add_picture(f"{loc}/{respic}",width=docx.shared.Cm(4))
   doc.add_picture(f"{loc}/{likpic}",width=docx.shared.Cm(4))

   doc.add_heading("Config",1)
   with open(f"{loc}/{conf}","r") as fil:
      cnf = fil.readlines()
   for line in cnf:
      if not line.lstrip()[0] == "#":
         doc.add_paragraph(line,style="Body Text")

   return doc

if __name__ == "__main__":

    parser = argparse.ArgumentParser(description="Script to make reports from OnOff-Analsis script")
    parser.add_argument("results_loc",
                       type=str,
                       help="Location of the analysis result")
    parser.add_argument("report_name",
                       type=str,
                       help="Name of file to save to")
    args = parser.parse_args()

    doc = make_report(args.results_loc)

    print(f"Saving result to {args.report_name}.docx")
    doc.save(f"{args.report_name}.docx")
